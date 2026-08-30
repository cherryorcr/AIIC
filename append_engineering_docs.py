from pathlib import Path
import sys

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from build_project_plan import (
    BLUE,
    NAVY,
    MUTED,
    LIGHT_BLUE,
    LIGHT_GOLD,
    GOLD,
    set_run_font,
    add_para,
    add_bullet,
    add_heading,
    add_callout,
    add_table,
)
from append_system_design import add_code_block, remove_end_marker


def add_labeled_para(doc, label, text, color=NAVY):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.08
    r = p.add_run(label + "：")
    set_run_font(r, size=10.5, color=color, bold=True)
    r = p.add_run(text)
    set_run_font(r, size=10.5, color="1F2937")
    return p


def build(input_path: Path, output_path: Path):
    doc = Document(input_path)
    remove_end_marker(doc)
    doc.add_page_break()

    add_para(doc, "软件工程文档补充", size=10, color=BLUE, bold=True, after=5)
    add_heading(doc, "第三部分：产品备忘录与工程详细设计补充", 1)
    add_para(doc, "本部分将仓库中的实现、测试和部署文件整理为可评审的产品与工程文档。实现现状使用 implemented 或 partial 标记；规划项使用 planned；真实用户和质量结论若尚未完成验证，使用 needs_validation。")
    add_callout(doc, "证据边界", "当前已纳入 1 条匿名真实调研摘录（R-001），但它只代表一个用户的需求信号，不代表总体比例。15 条问卷样例和早期题目样例仍为 synthetic_mock；本方案不将其表述为真实用户比例、真实面试频率或公司官方题库。", fill=LIGHT_GOLD, accent=GOLD)

    add_heading(doc, "27. Product Memo", 1)
    add_heading(doc, "27.1 目标用户与核心痛点", 2)
    add_para(doc, "首要用户是准备实习、校招或保研面试的本科生和研究生，重点覆盖软件开发、算法、数据、产品和科研方向。根据 R-001，用户可按组织方式分为两类：单面用户需要独立练习简历面、专业面、论文/科技面、实习面、行为面或英文面，并可能按学校/公司流程组合多个环节；群面用户希望与水平相近且时间重合的同学组成小组或圈子，成员能够互相看见表现并共同学习。")
    add_para(doc, "已确认的需求输入还包括：依据脱敏 JD、项目经历和技能获得岗位/题目匹配；简历面持续围绕简历深挖；自动录屏和自动转写用于复盘、证据提取和训练报告；技术、算法、压力、案例、科研和 HR 场景使用独立策略；算法题能够运行代码并获得复杂度反馈；业务服务、双 4090 本地模型和第三方强模型通过网关解耦。")
    add_para(doc, "R-001 只是一条真实调研信号，不能推导总体比例、普遍频率或市场结论。待验证假设包括：多环节组合是否是多数流程；用户是否愿意提供空闲时间和能力标签；录屏/转写是否被多数用户视为基础能力；群面等待池的隐私、可见范围和等待时长如何设置；用户最痛的是找题、判断回答质量还是缺少连续追问。真实访谈计划为 5–8 人、问卷 20–30 份，结果应回填 docs/product-memo.md 的访谈表。")
    add_table(doc, ["证据层", "当前状态", "允许的结论"], [
        ("项目发起人的需求输入", "observed_requirement", "确认产品方向和功能偏好，不等同用户研究"),
        ("匿名真实调研摘录 R-001", "observed_interview", "支持单面/群面分类、单面环节组合、简历深挖、群面组队、录屏转写等需求信号；不代表总体比例"),
        ("mock-survey-responses.csv", "synthetic_mock", "验证分析流程和演示页面，不代表真实比例"),
        ("扩大后的真实访谈/问卷", "needs_validation", "完成后才能形成稳定的场景、痛点和优先级结论"),
    ], [2.05, 1.55, 2.9], font_size=8.3)

    add_heading(doc, "27.1.1 真实调研摘录与初步结论", 3)
    add_table(doc, ["记录", "用户表达/场景", "需求含义", "置信边界"], [
        ("R-001-A", "第一种是群面和单面", "产品一级导航先区分群面/单面", "单个样本"),
        ("R-001-B", "单面包括简历、专业、论文/科技、实习、行为、英文；学校面试可能包含多个环节", "支持单面环节选择与组合", "需要更多学校/公司样本"),
        ("R-001-C", "今天想练简历面，希望持续围绕简历把可挖掘内容问到", "简历深挖、经历证据链、已问/未问覆盖", "需验证全量深挖与限时训练取舍"),
        ("R-001-D", "难以找到水平相近且刚好同一时间有空的多人；希望自由组队/组圈，成员视频互相可见", "能力相似度 + 时间重合 + 人数 + 可见范围匹配", "需调查隐私、等待池和缺席处理"),
        ("R-001-E", "自动录屏和自动转写是基础功能", "录屏/ASR 用于复盘、证据提取和报告", "当前先做接口与数据契约"),
    ], [0.75, 2.6, 2.15, 1.0], font_size=7.6)
    add_para(doc, "初步产品判断：群面不是简单的多人聊天页面，而是带能力、空闲时间、人数和可见范围约束的协同训练场景；单面不是一个统一模式，而是可组合的专业环节；录屏和转写的主要价值是复盘与证据提取，实时数字人暂不能替代真实群面匹配。")
    add_labeled_para(doc, "优先级影响", "P0 保持当前文字单面闭环、算法执行、岗位/题目匹配、反馈、追问和重答；P1 增加简历深挖、可组合单面流程、群面匹配等待池、群面房间与录屏/转写接口；P2 再实现浏览器录屏、ASR、多人实时媒体复盘和可控的视频可见范围。")

    add_heading(doc, "27.2 产品设计与刻意不做的功能", 2)
    add_table(doc, ["核心功能", "用户价值", "状态"], [
        ("资料上传、解析、人工确认", "减少录入，避免 AI 抽取结果直接成为事实", "implemented"),
        ("岗位/技能/题目匹配", "展示技能覆盖、缺口、来源和匹配理由", "implemented"),
        ("八类场景面试", "为每类流程切换 Prompt 和 Rubric", "implemented"),
        ("原句证据反馈、追问、重答", "把泛化反馈变成可定位的下一步动作", "implemented + fallback"),
        ("Python 算法题沙箱", "验证测试用例、复杂度和代码质量", "partial，生产需独立 runner"),
        ("Model Router", "在本地模型、第三方模型和规则 fallback 间切换", "implemented"),
    ], [2.1, 3.35, 1.05], font_size=8.5)
    add_para(doc, "刻意暂不做：实时语音、视频、数字人；自训练基础模型；Kubernetes；支付、社区和商业化；未经授权的公司内部面经抓取。原因是这些功能会引入媒体链路、隐私、版权、运维或商业范围，不能直接证明当前核心闭环价值。")

    add_heading(doc, "27.3 版本迭代记录", 2)
    add_table(doc, ["版本/问题", "改动", "原因"], [
        ("V0：全功能 AI 面试官", "收敛为文字闭环、算法题和场景策略", "一天内先验证岗位/证据到反馈的核心链路"),
        ("全远程 API 成本和故障", "引入双 4090 本地模型、统一 Model Router 和强模型 fallback", "降低成本和供应商耦合，支持降级"),
        ("真实题库版权和频率不可证", "来源台账、许可核验、公开题转写、synthetic_mock 标记", "可公开演示且不误导"),
        ("一次性参考答案过泛", "场景 Rubric、原句证据、追问和 revision_of 重答", "使反馈可执行且可持续改进"),
        ("直接执行用户代码有风险", "AST、audit hook、资源限制和超时；生产迁移独立 runner", "算法面可用，同时明确安全边界"),
    ], [1.85, 3.2, 1.45], font_size=8.3)

    add_heading(doc, "27.4 如果再给一周", 2)
    for t in [
        "第 1–2 天：完成 5–8 个半结构化访谈和 20–30 份匿名问卷，做主题编码并回填 Product Memo。",
        "第 3 天：建立 5 组脱敏 JD、5 组用户背景、20–30 道有许可/自编题和人工金标准。",
        "第 4 天：增加反馈人工复核、事实错误标记、证据缺失标记和 Prompt/Rubric 版本化。",
        "第 5 天：对本地模型、第三方强模型和规则 fallback 比较质量、p50/p95 延迟、token 和成本。",
        "第 6 天：加入 Redis/worker 和独立 sandbox runner，完成 GPU 网关鉴权、备份恢复和故障演练。",
        "第 7 天：让真实用户完成训练，修复最高影响的三个问题，更新 README、Demo、Memo 和验收结果。",
    ]:
        add_bullet(doc, t)

    add_heading(doc, "27.5 AI 工具使用", 2)
    add_table(doc, ["工具/能力", "使用环节", "人工责任"], [
        ("ChatGPT/Codex 类强模型", "需求梳理、文档草拟、复杂反馈和代码协作", "核对实现、事实、范围、隐私和来源"),
        ("本地 Qwen/vLLM", "JD/简历抽取、轻量出题、结构化 fallback", "校验 schema、质量、显存和延迟"),
        ("Embedding/Reranker", "技能与题目语义召回", "建立人工评测集，保留规则 fallback"),
        ("GitHub API/网页阅读", "来源、版本和 MIT 许可核验", "判断转写和再分发边界"),
        ("文档/测试工具", "Markdown、DOCX、结构和可访问性 QA", "最终审阅、运行验证和发布"),
    ], [1.8, 2.7, 2.0], font_size=8.4)

    add_heading(doc, "28. 用例设计补充", 1)
    add_para(doc, "正式用例见 docs/use-case-design.md。本节保留评审所需的主参与者、业务边界和关键规则摘要。")
    add_code_block(doc, """候选人 -> UC-01 工作区 -> UC-03 背景 -> UC-04 资料确认
        -> UC-05 岗位/题目匹配 -> UC-06 创建会话
        -> UC-07 文字面试/追问 -> UC-09 重答
        -> UC-08 算法代码（算法面） -> UC-10 报告
        -> UC-11 题库收藏 -> UC-12 历史准备度

知识库管理员 -> UC-13 题库、来源和许可证维护
系统管理员   -> UC-14 模型路由、健康、成本和错误

统一规则：服务端会话推导 user_id；反馈只能引用原回答；
题库必须带来源/许可；模型失败和沙箱失败不能丢输入。""")
    add_table(doc, ["用例", "主成功结果", "关键异常"], [
        ("UC-04 资料确认", "用户确认后的 Profile/Job 可用于匹配", "文件不支持、过大或解析失败，保留失败状态"),
        ("UC-05 匹配", "返回技能覆盖、缺口、题目、来源和置信度", "无结果时使用有许可/自编基础题"),
        ("UC-07 面试回答", "Turn + Feedback 持久化并推进追问", "模型超时、非法 JSON 时降级"),
        ("UC-08 算法执行", "返回通过数、错误、耗时和安全状态", "rejected/timeout/error/disabled"),
        ("UC-10 报告", "完成会话并保存报告快照", "零回答会话不生成零分报告"),
    ], [1.65, 3.15, 1.7], font_size=8.5)

    add_heading(doc, "29. 系统概要设计补充", 1)
    add_para(doc, "系统采用前后端分离和双主机部署：React/Nginx 负责展示，FastAPI 负责业务 API 和会话编排，PostgreSQL 负责生产持久化，GraphRAG 负责可解释召回，Model Router 连接双 4090 网关和第三方强模型，算法代码进入受限 runner。当前 CPU Compose 已包含 frontend、backend 和 PostgreSQL；Redis/worker 与独立 runner 是并发和公网开放前的目标拆分。")
    add_code_block(doc, """Browser
  -> Nginx frontend (/)
  -> Nginx proxy (/api) -> FastAPI backend
       -> Database (SQLite dev / PostgreSQL prod)
       -> GraphRAG (relation filter + vector/lexical rank)
       -> Model Router
            -> private GPU Model Gateway (Qwen/vLLM)
            -> third-party strong model relay
       -> SandboxService (Python challenge)
""")
    add_table(doc, ["质量属性", "设计措施"], [
        ("可解释性", "matched skills、source refs、confidence、evidence quotes"),
        ("降级性", "local/strong/规则 fallback，健康状态不阻塞 API 启动"),
        ("隐私", "HttpOnly Cookie、服务端 owner 校验、材料确认后落规范表"),
        ("可扩展", "场景 policy/Rubric、Model Router adapter、Database 边界"),
        ("可运营", "provider、model、latency、tokens、cost、fallback 遥测"),
    ], [1.55, 4.95], font_size=8.7)

    add_heading(doc, "30. 前后端详细设计", 1)
    add_heading(doc, "30.1 前端", 2)
    add_table(doc, ["路由", "职责", "主要 API"], [
        ("/", "准备度、最近训练、建议", "/workspace/overview、/history"),
        ("/practice", "场景、问题、回答、追问、算法代码", "/scenes、/sessions、/turns、/algorithm/run"),
        ("/match", "岗位、技能覆盖、缺口、题目", "/jobs、/matches"),
        ("/materials", "上传、解析、确认简历/JD", "/documents/*、/profile、/jobs"),
        ("/questions", "公开题库、筛选、收藏、加入训练", "/questions、/favorites"),
        ("/reports", "历史、摘要、反馈详情", "/history、/reports、/sessions/{id}"),
    ], [1.25, 3.0, 2.25], font_size=8.3)
    add_para(doc, "PracticePage 将状态分为 Session、Answer、Feedback 和 UI 四类。会话主状态为 idle → starting → questioning → submitting/running_code → feedback → retry/supplement/next → completed。api.ts 统一处理 Cookie、超时、瞬时错误重试和 POST 禁止盲重试。")
    add_heading(doc, "30.2 后端", 2)
    add_table(doc, ["层", "职责", "当前文件"], [
        ("API 边界", "路由、鉴权、Pydantic 校验、错误转换", "backend/app/main.py、schemas/models.py"),
        ("应用服务", "会话、匹配、资料、报告、知识库", "services/interview.py、main.py"),
        ("领域能力", "GraphRAG、Prompt、Model Router、Sandbox", "services/rag.py、prompts.py、model_router.py、sandbox.py"),
        ("基础设施", "事务、迁移、SQLite/PostgreSQL 兼容", "storage/db.py、migrations/"),
    ], [1.35, 3.4, 1.75], font_size=8.3)
    add_para(doc, "回答流程先保存 Turn，再调用模型评价；反馈 schema 校验通过后保存 Feedback；随后推进当前问题并 upsert 报告快照。重答通过 revision_of 建立父子 Turn，不移动已经进入追问流程的会话游标。")

    add_heading(doc, "31. GraphRAG、模型路由与算法沙箱详细设计", 1)
    add_heading(doc, "31.1 GraphRAG", 2)
    add_code_block(doc, """role/JD/profile
  -> extract_skills + normalize_skill
  -> graph filter: process_type + role + difficulty + skill edges
  -> lexical/hash-vector similarity
  -> semantic dedupe + source-confidence rerank
  -> matched_skills + questions + source_refs + confidence""")
    add_para(doc, "当前 GraphRAG 使用关系表、关键词/确定性哈希向量和语义去重；后续可在同一 Repository 边界内迁移 PostgreSQL + pgvector 或真正的 embedding/reranker，不改变 API 输出。")
    add_heading(doc, "31.2 Model Router", 2)
    add_table(doc, ["任务", "默认路由", "降级"], [
        ("抽取/技能归一化", "本地模型或规则", "强模型或确定性字段解析"),
        ("题目召回/重排", "本地 embedding/rerank", "关键词和技能关系"),
        ("个性化出题/压力追问", "第三方强模型", "本地模型/题库原题"),
        ("结构化评分", "本地模型 + Rubric", "强模型后规则评分"),
        ("代码执行", "本地 Sandbox", "不交给模型执行"),
    ], [1.7, 2.35, 2.45], font_size=8.5)
    add_para(doc, "每次调用记录 provider、model、task、route、latency、tokens、status、error_name、fallback_reason 和 prompt_version。对 JSON 输出先校验，再尝试一次 repair，失败时使用 InterviewService 的确定性反馈。")
    add_heading(doc, "31.3 算法沙箱", 2)
    add_para(doc, "当前实现使用 AST 拒绝危险导入/调用/反射，audit hook 拒绝 socket、子进程、系统命令和动态库加载，临时目录运行固定 harness，并限制 wall timeout、CPU、内存、进程、文件大小、文件描述符和输出。该实现适合挑战演示；公网开放前必须迁移独立非 root runner，启用禁网、只读文件系统、seccomp/AppArmor 或 nsjail，不挂载宿主目录和 Docker Socket。")

    add_heading(doc, "32. API 与数据契约", 1)
    add_para(doc, "完整接口见 docs/api-design.md。新开发统一使用 /api/v1；/api/interview/start 和 /api/interview/answer 仅作为历史兼容别名。")
    add_table(doc, ["接口", "输入", "输出/约束"], [
        ("POST /sessions", "mode、role、job_text、user_profile、difficulty", "session_id、首题、匹配技能、来源"),
        ("POST /sessions/{id}/turns", "question_id、answer_text/code、revision_of", "feedback、next_question、algorithm_result"),
        ("POST /matches", "岗位、JD、用户画像", "match_score、required/matched/gap skills、题目"),
        ("POST /documents/parse", "multipart file、kind", "解析草稿，必须 confirm 才写规范资料"),
        ("POST /sessions/{id}/complete", "无", "completed 状态和报告快照"),
        ("GET /model/health", "无", "provider 状态、延迟、错误、fallback"),
    ], [2.3, 2.3, 1.9], font_size=8.2)
    add_callout(doc, "数据边界", "Session、Turn、Feedback、Report、UserProfile、Job、Question、Skill、Source 和 ModelInvocation 通过 Database 持久化。所有用户资源先做 owner 校验；题目来源保留 URL、许可证、版本、hash、置信度和脱敏状态。", fill=LIGHT_BLUE, accent=BLUE)

    add_heading(doc, "33. 测试计划、验收与需求追踪", 1)
    add_para(doc, "docs/test-plan.md 定义单元、API 集成、前端契约、端到端、安全、部署和质量评测层次；docs/requirements-traceability.md 将 FR/NFR 需求关联到用例、代码和测试。当前 test_smoke.py 已覆盖会话、重答、匹配、题库、沙箱、认证、资料、报告、PostgreSQL 迁移和 Model Router fallback。")
    add_table(doc, ["验收项", "目标", "当前状态"], [
        ("核心文字闭环", "资料/岗位 → 匹配 → 回答 → 反馈 → 追问/报告", "implemented"),
        ("场景覆盖", "技术、算法、行为、压力、案例、科研、HR、群面", "implemented"),
        ("算法安全", "正确、失败、超时、违规均有结构化结果", "partial，生产 runner 待拆分"),
        ("来源可信", "题目显示来源、许可、版本、置信度", "implemented"),
        ("质量效果", "匹配和反馈人工金标准、真实用户满意度", "needs_validation"),
        ("自动化回归", "pytest 全部通过", "待安装 requirements-dev 后执行"),
    ], [1.7, 3.7, 1.1], font_size=8.4)

    add_heading(doc, "34. ADR 与工程治理", 1)
    add_table(doc, ["决策", "结论", "保留的代价/风险"], [
        ("ADR-001 部署", "Docker Compose 双主机，不上 Kubernetes", "需手工维护两台主机和发布流程"),
        ("ADR-002 模型", "业务只依赖 Model Router", "需维护 provider 兼容和标准 schema"),
        ("ADR-003 GraphRAG", "关系表 + 轻量向量，后续 pgvector", "MVP 召回质量需要人工评测"),
        ("ADR-004 数据", "许可核验、来源台账、自行转写", "不能宣称官方频率，需定期复核"),
        ("ADR-005 交互", "当前先做文字，多模态只留接口", "暂不能评价语音/视频临场感"),
        ("ADR-006 算法", "模型不执行代码，受限 runner 执行", "需要持续维护隔离策略"),
        ("ADR-007 存储", "开发 SQLite，生产 PostgreSQL", "需持续验证双数据库兼容"),
    ], [1.55, 3.3, 1.65], font_size=8.0)
    add_para(doc, "文档维护规则：新增 API 先改契约；新增场景同步 policy、Rubric、用例和测试；新增题目先登记来源和许可证；产品结论必须注明 evidence level；每次发布记录版本、日期、变更原因和未解决风险。")

    add_heading(doc, "35. 文档补齐结论", 1)
    add_callout(doc, "交付结论", "当前项目已形成从产品需求、用例、概要架构、前后端详细设计、GraphRAG、模型路由、算法沙箱、API 契约、数据库、测试、部署到 ADR 的完整文档链路。唯一不能用文档替代的部分是真实用户研究和人工质量评测；这些应在下一轮访谈、问卷和固定评测集完成后回填 Product Memo 与追踪矩阵。", fill=LIGHT_BLUE, accent=BLUE)
    add_para(doc, "—— 文档结束 ——", size=9, color=MUTED, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, before=16, after=2)
    doc.save(output_path)
    print(output_path)


if __name__ == "__main__":
    if len(sys.argv) != 3:
        raise SystemExit("usage: append_engineering_docs.py INPUT.docx OUTPUT.docx")
    build(Path(sys.argv[1]), Path(sys.argv[2]))
